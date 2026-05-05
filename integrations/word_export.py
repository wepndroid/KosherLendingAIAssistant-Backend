from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from typing import Iterable


NAVY = RGBColor(0x1A, 0x3C, 0x6E)
GOLD = RGBColor(0xC9, 0x97, 0x3A)


def render_content_pack(items: Iterable[dict], title: str, compliance_footer: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.color.rgb = NAVY

    for item in items:
        doc.add_heading(item.get("topic") or "Untitled", level=2)

        meta = doc.add_paragraph()
        meta.add_run(
            f"Pillar: {item.get('pillar','-')}  ·  Platform: {item.get('platform','-')}  ·  Duration: {item.get('duration','-')}"
        ).italic = True

        _section(doc, "Hook", item.get("hook"))
        _section(doc, "On-screen text", item.get("on_screen"))
        _section(doc, "Script", item.get("script"))
        _section(doc, "Production brief", item.get("production_brief"))

        doc.add_paragraph().add_run("Captions").bold = True
        for label, key in [
            ("TikTok", "caption_tiktok"),
            ("Instagram", "caption_instagram"),
            ("LinkedIn", "caption_linkedin"),
            ("Facebook", "caption_facebook"),
            ("X/Twitter", "caption_x"),
        ]:
            v = item.get(key) or item.get("caption")
            if v:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(v)

        cta_strategy = ((item.get("validations") or {}).get("cta_strategy") if isinstance(item.get("validations"), dict) else None)
        _section(doc, "CTA strategy", cta_strategy)
        _section(doc, "Spoken CTA", item.get("cta"))
        _section(doc, "DM keyword → deliverable", f"{item.get('dm_keyword','-')} → {item.get('deliverable','-')}")

        tags = item.get("hashtags") or []
        if tags:
            _section(doc, "Hashtags", " ".join(tags))

        src = doc.add_paragraph()
        src.add_run("Source: ").bold = True
        src.add_run(f"{item.get('source_book','-')} — {item.get('source_framework','-')}")

        compliance = doc.add_paragraph(compliance_footer)
        for run in compliance.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = GOLD

        doc.add_paragraph("─" * 40)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _section(doc: Document, label: str, value):
    if not value:
        return
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(str(value))
