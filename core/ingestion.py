"""Document ingestion: parse → semantic-chunk → embed → store.

Designed for files up to ~3 GB. Two entry points:
- `ingest_document(document_id, filename, data: bytes)` — backwards-compatible small-file path
- `ingest_streaming(document_id, filename, file_obj)` — streaming path; reads page-by-page / paragraph-by-paragraph so memory stays bounded regardless of file size

Chunking is boundary-aware: paragraph → sentence → word fallback, with section headings
attached to each chunk's metadata so citations stay precise.
"""
from __future__ import annotations
import io
import re
from pathlib import Path
from typing import Iterable, Iterator
from docx import Document as DocxDocument
from pypdf import PdfReader
from db.supabase_client import supabase
from integrations.openai_client import embed
from integrations.claude_client import summarize


# Target chunk size in *words* (close enough to tokens for English prose; ~600 tokens)
CHUNK_TARGET_WORDS = 600
CHUNK_MAX_WORDS = 800
CHUNK_OVERLAP_WORDS = 80

# Heuristics for section heading detection
_HEADING_RE = re.compile(
    r"^(?:"
    r"(?:CHAPTER|PART|SECTION|APPENDIX|PROLOGUE|EPILOGUE|INTRODUCTION|CONCLUSION)\b.*"
    r"|\s*\d+(?:\.\d+)*\s+[A-Z][^.!?]{3,80}"   # "1.2 Loss Aversion"
    r"|[A-Z][A-Z0-9 ,'\-:]{3,80}"               # ALL-CAPS line
    r"|[A-Z][^.!?]{0,80}:"                      # "Key insight:"
    r")$"
)


def parse_text(filename: str, data: bytes) -> str:
    """Backwards-compatible: returns whole text as one string. OK for files under ~50 MB."""
    return "\n".join(b.text for b in parse_blocks(filename, io.BytesIO(data)))


# ─── Streaming parser ───────────────────────────────────────────

class Block:
    """A chunk of text with structural metadata (section, page, paragraph index)."""
    __slots__ = ("text", "section", "page", "paragraph_index")

    def __init__(self, text: str, *, section: str | None = None, page: int | None = None, paragraph_index: int | None = None):
        self.text = text
        self.section = section
        self.page = page
        self.paragraph_index = paragraph_index


def parse_blocks(filename: str, file_obj) -> Iterator[Block]:
    """Yield (text, metadata) blocks from a file-like object. Streams page-by-page for PDFs and paragraph-by-paragraph for DOCX."""
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        yield from _parse_docx_blocks(file_obj)
    elif ext == ".pdf":
        yield from _parse_pdf_blocks(file_obj)
    elif ext in {".txt", ".md"}:
        yield from _parse_text_blocks(file_obj)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_docx_blocks(file_obj) -> Iterator[Block]:
    doc = DocxDocument(file_obj)
    current_section: str | None = None
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text:
            continue
        # Use the docx style as a hint for section headings
        style = (getattr(getattr(p, "style", None), "name", "") or "").lower()
        if style.startswith("heading") or _is_heading(text):
            current_section = text[:120]
            continue
        yield Block(text, section=current_section, paragraph_index=i)
    # Tables — flatten into paragraph-style blocks
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                t = (cell.text or "").strip()
                if t:
                    yield Block(t, section=current_section or f"Table {ti+1}", paragraph_index=ri * 1000 + ci)


def _parse_pdf_blocks(file_obj) -> Iterator[Block]:
    reader = PdfReader(file_obj)
    current_section: str | None = None
    for pi, page in enumerate(reader.pages, 1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if not page_text.strip():
            continue
        # Split page into paragraphs by blank lines
        for para_idx, para in enumerate(re.split(r"\n\s*\n", page_text)):
            text = re.sub(r"\s+", " ", para).strip()
            if not text:
                continue
            if _is_heading(text):
                current_section = text[:120]
                continue
            yield Block(text, section=current_section, page=pi, paragraph_index=para_idx)


def _parse_text_blocks(file_obj) -> Iterator[Block]:
    # Read in 256 KB windows, accumulating to paragraph boundaries
    buf = ""
    current_section: str | None = None
    para_idx = 0
    while True:
        data = file_obj.read(262144)
        if not data:
            break
        if isinstance(data, bytes):
            data = data.decode("utf-8", errors="ignore")
        buf += data
        # Flush all complete paragraphs
        while True:
            split = re.search(r"\n\s*\n", buf)
            if not split:
                break
            chunk_text = buf[: split.start()].strip()
            buf = buf[split.end():]
            if chunk_text:
                if _is_heading(chunk_text):
                    current_section = chunk_text[:120]
                else:
                    yield Block(chunk_text, section=current_section, paragraph_index=para_idx)
                    para_idx += 1
    if buf.strip():
        yield Block(buf.strip(), section=current_section, paragraph_index=para_idx)


def _is_heading(text: str) -> bool:
    if len(text) > 120:
        return False
    if not text:
        return False
    return bool(_HEADING_RE.match(text))


# ─── Boundary-aware chunker ─────────────────────────────────────

def semantic_chunk_iter(
    blocks: Iterable[Block],
    *,
    target: int = CHUNK_TARGET_WORDS,
    max_size: int = CHUNK_MAX_WORDS,
    overlap: int = CHUNK_OVERLAP_WORDS,
) -> Iterator[dict]:
    """Yield chunks while streaming blocks, preserving boundaries + metadata."""
    current_words: list[str] = []
    current_meta: dict = {}

    def flush() -> dict | None:
        if not current_words:
            return None
        text = " ".join(current_words).strip()
        if text:
            return {"text": text, "metadata": dict(current_meta)}
        return None

    for blk in blocks:
        words = blk.text.split()
        # If the block changes section relative to the current chunk, flush first so
        # each chunk's metadata stays accurate for citation
        if current_words and current_meta.get("section") != blk.section:
            item = flush()
            if item:
                yield item
            current_words = []
            current_meta = {}

        # If a single block is bigger than max_size, sentence-split it
        sub_blocks = _split_oversize(words, max_size) if len(words) > max_size else [words]

        for sub in sub_blocks:
            # If adding this sub would exceed max_size, flush current with overlap
            if current_words and len(current_words) + len(sub) > max_size:
                item = flush()
                if item:
                    yield item
                current_words = current_words[-overlap:] if overlap > 0 else []
                current_meta = {**current_meta, "carried_overlap": True}

            # If chunk is currently empty, seed metadata from this block
            if not current_words:
                current_meta = {
                    "section": blk.section,
                    "page": blk.page,
                    "paragraph_index": blk.paragraph_index,
                }
            current_words.extend(sub)

            # If we're past target and at a paragraph boundary (end of this block), flush
            if len(current_words) >= target:
                item = flush()
                if item:
                    yield item
                current_words = current_words[-overlap:] if overlap > 0 else []
                current_meta = {**current_meta, "carried_overlap": True}

    item = flush()
    if item:
        yield item


def semantic_chunk(
    blocks: Iterable[Block],
    *,
    target: int = CHUNK_TARGET_WORDS,
    max_size: int = CHUNK_MAX_WORDS,
    overlap: int = CHUNK_OVERLAP_WORDS,
) -> list[dict]:
    """Pack blocks into chunks that respect paragraph/sentence boundaries.

    Output: list of {"text": str, "metadata": {...}}
    """
    return list(semantic_chunk_iter(blocks, target=target, max_size=max_size, overlap=overlap))


def _split_oversize(words: list[str], max_size: int) -> list[list[str]]:
    """Split an oversized block on sentence boundaries first; word boundary as last resort."""
    text = " ".join(words)
    sentences = re.split(r"(?<=[.!?])\s+", text)
    out: list[list[str]] = []
    cur: list[str] = []
    for s in sentences:
        sw = s.split()
        if len(cur) + len(sw) > max_size and cur:
            out.append(cur)
            cur = []
        if len(sw) > max_size:
            # Sentence itself too long — break by word
            for i in range(0, len(sw), max_size):
                out.append(sw[i : i + max_size])
            continue
        cur.extend(sw)
    if cur:
        out.append(cur)
    return out


# ─── Public API ─────────────────────────────────────────────────

def chunk(text: str, size: int = CHUNK_TARGET_WORDS, overlap: int = CHUNK_OVERLAP_WORDS) -> list[str]:
    """Backwards-compat wrapper: returns plain strings (no metadata).

    New code should call `semantic_chunk(parse_blocks(...))` for richer metadata.
    """
    blocks = [Block(p, paragraph_index=i) for i, p in enumerate(re.split(r"\n\s*\n", text)) if p.strip()]
    if not blocks:
        return []
    rich = semantic_chunk(blocks, target=size, overlap=overlap)
    return [c["text"] for c in rich]


def ingest_document(*, document_id: str, filename: str, data: bytes) -> dict:
    """Small-file path: parse all into RAM, chunk, embed, store. OK for files <~ 50 MB."""
    return _run_pipeline(document_id, filename, io.BytesIO(data))


def ingest_streaming(*, document_id: str, filename: str, file_obj) -> dict:
    """Streaming path: file_obj is a seekable file-like; pdf/docx readers parse page/paragraph at a time."""
    return _run_pipeline(document_id, filename, file_obj)


def _run_pipeline(document_id: str, filename: str, file_obj) -> dict:
    db = supabase()
    db.table("knowledge_documents").update({"status": "Processing"}).eq("id", document_id).execute()
    chunk_index = 0
    batch_chunks: list[dict] = []
    batch_texts: list[str] = []
    summary_head_parts: list[str] = []
    summary_head_chars = 0
    summary_head_limit = 30_000

    def flush_batch() -> int:
        nonlocal chunk_index, batch_chunks, batch_texts, summary_head_chars
        if not batch_chunks:
            return 0

        vectors = embed(batch_texts)
        rows = []
        for c, v in zip(batch_chunks, vectors):
            txt = c["text"]
            if summary_head_chars < summary_head_limit:
                take = min(summary_head_limit - summary_head_chars, len(txt))
                summary_head_parts.append(txt[:take])
                summary_head_chars += take
            rows.append(
                {
                    "document_id": document_id,
                    "chunk_index": chunk_index,
                    "chunk_text": txt,
                    "embedding": v,
                    "metadata": {"source": filename, **(c.get("metadata") or {})},
                }
            )
            chunk_index += 1

        db.table("knowledge_chunks").insert(rows).execute()
        count = len(rows)
        batch_chunks = []
        batch_texts = []
        return count

    total_chunks = 0
    had_text = False
    for c in semantic_chunk_iter(parse_blocks(filename, file_obj)):
        had_text = True
        batch_chunks.append(c)
        batch_texts.append(c["text"])
        if len(batch_chunks) >= 100:
            total_chunks += flush_batch()
            if total_chunks % 500 == 0:
                db.table("knowledge_documents").update({"total_chunks": total_chunks}).eq("id", document_id).execute()

    total_chunks += flush_batch()

    if not had_text or total_chunks == 0:
        db.table("knowledge_documents").update({"status": "Failed", "summary": "No text extracted"}).eq("id", document_id).execute()
        return {"status": "Failed", "reason": "No text extracted"}

    summary = ""
    try:
        summary = summarize(" ".join(summary_head_parts)[:summary_head_limit])
    except Exception:
        pass

    db.table("knowledge_documents").update(
        {
            "status": "Indexed",
            "total_chunks": total_chunks,
            "summary": summary,
            "indexed_at": "now()",
        }
    ).eq("id", document_id).execute()

    db.table("activity_log").insert(
        {"text": f"{filename} indexed ({total_chunks} chunks)", "icon": "check"}
    ).execute()

    return {"status": "Indexed", "chunks": total_chunks, "summary": summary}
